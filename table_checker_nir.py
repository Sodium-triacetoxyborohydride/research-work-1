from docx import Document


def check_tables(doc):
    tables = doc.tables
    errors = []

    for i, table in enumerate(tables):
        # Проверка на наличие заголовка таблицы
        if not table.rows[0].cells[0].text.startswith("Таблица"):
            errors.append(f"Таблица {i + 1}: Отсутствует заголовок или он неправильно оформлен.")

        # Проверка правильной нумерации
        if not table.rows[0].cells[0].text.startswith(f"Таблица {i + 1}"):
            errors.append(f"Таблица {i + 1}: Неправильная нумерация таблицы.")

    return errors


def check_table_format(doc):
    errors = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Проверка правильности форматирования текста в ячейках
                if cell.text:
                    for run in cell.paragraphs[0].runs:
                        if run.font.name != 'Times New Roman' or run.font.size.pt != 12:
                            errors.append(
                                f"Ячейка с текстом '{cell.text[:30]}...' не соответствует шрифту Times New Roman 12pt.")

    return errors
